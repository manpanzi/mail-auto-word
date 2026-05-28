# -*- coding: utf-8 -*-
"""
mail_auto_word.py 鈥?閭欢鐩戞帶 & 鑷姩鐢熸垚 Word 骞跺洖澶?=====================================================
鍔熻兘锛?  1. 閫氳繃 IMAP 鐩戞帶鎸囧畾閭锛屽畾鏃舵鏌ユ柊閭欢
  2. 璇嗗埆涓婚涓虹┖ + 甯?Excel 闄勪欢锛?xlsx/.xls锛夌殑閭欢
  3. 瑙ｆ瀽閭欢闄勪欢涓殑 Excel 琛ㄦ牸锛?xlsx锛夛紝鎴栨鏂囦腑鐨?CSV 鏁版嵁
  4. 鑷姩璇嗗埆鏁版嵁绫诲瀷锛堝嚭搴撻€氱煡鍗?/ 鎻愯揣濮旀墭鍑斤級锛岀敓鎴?Word 鏂囨。
  5. 灏嗙敓鎴愮殑 .docx 浣滀负闄勪欢鍥炲缁欏彂浠朵汉

鏀寔閭锛歈Q 閭锛園qq.com锛夈€?63 閭锛園163.com锛?
鐢ㄦ硶锛?  python mail_auto_word.py          # 鍓嶅彴杩愯锛孋trl+C 鍋滄
  python mail_auto_word.py --once   # 鍙鏌ヤ竴娆″氨閫€鍑?"""

import os
import re
import sys
import csv
import json
import time
import logging
import zipfile
import argparse
import tempfile
import threading
import email.policy
from io import StringIO
from datetime import datetime
from pathlib import Path

# ============================================================
#  渚濊禆妫€鏌ワ紙棣栨杩愯鎻愮ず瀹夎锛?# ============================================================
MISSING = []
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    MISSING.append("python-docx")

try:
    from openpyxl import Workbook, load_workbook
except ImportError:
    MISSING.append("openpyxl")

if MISSING:
    print("Missing deps: " + ", ".join(MISSING))
    print("Run: pip install " + " ".join(MISSING))
    sys.exit(1)

# ============================================================
#  閰嶇疆鍖?鈥?鎸夐渶淇敼
# ============================================================

# --- 閭璐﹀彿 ---
# QQ 閭 / 163 閭鍧囧彲锛屼娇鐢?鎺堟潈鐮侊紙闈炵櫥褰曞瘑鐮侊級鐧诲綍
# 濡備綍鑾峰彇鎺堟潈鐮侊細
#   QQ 閭锛氳缃?鈫?璐︽埛 鈫?POP3/IMAP/SMTP 鏈嶅姟 鈫?寮€鍚?IMAP/SMTP 鈫?鐢熸垚鎺堟潈鐮?#   163 閭锛氳缃?鈫?POP3/SMTP/IMAP 鈫?寮€鍚?IMAP/SMTP 鈫?鏂板鎺堟潈鐮?# --- 閭璐﹀彿锛堜紭鍏堜粠鐜鍙橀噺璇诲彇锛岄€傜敤浜庝簯绔儴缃诧級---
ACCOUNT = {
    "email":    os.environ.get("MAIL_USER", "357818590@qq.com"),
    "password": os.environ.get("MAIL_PASS", "ukkdcvbewfesbhdc"),
}

# --- IMAP 鏀朵欢鏈嶅姟鍣紙鏍规嵁閭鑷姩閫夋嫨锛屼竴鑸棤闇€淇敼锛?--
# QQ:  imap.qq.com,  993, SSL
# 163: imap.163.com, 993, SSL
IMAP_CONFIG = {
    "qq.com":  {"host": "imap.qq.com",  "port": 993},
    "163.com": {"host": "imap.163.com", "port": 993},
}

# --- SMTP 鍙戜欢鏈嶅姟鍣?---
# QQ:  smtp.qq.com,  465, SSL
# 163: smtp.163.com, 465, SSL
SMTP_CONFIG = {
    "qq.com":  {"host": "smtp.qq.com",  "port": 465},
    "163.com": {"host": "smtp.163.com", "port": 465},
}

# --- 鐩戞帶璁剧疆 ---
CHECK_INTERVAL = 60          # 妫€鏌ラ棿闅旓紙绉掞級锛屽缓璁?鈮?30
MARK_AS_READ   = True        # 澶勭悊瀹屾垚鍚庢槸鍚︽爣璁颁负宸茶
TRACK_FILE     = "processed_uids.json"  # 璁板綍宸插鐞嗛偖浠剁殑 UID锛岄伩鍏嶉噸澶嶅鐞?
# --- 閭欢璇嗗埆 ---
SUBJECT_KEYWORDS = ["GENWORD"]  # 宸插簾寮冿紝鐜板湪閫氳繃 Excel 闄勪欢鍖归厤

# --- 鏃ュ織 ---
LOG_FILE = "mail_auto_word.log"

# ============================================================
#  鏃ュ織閰嶇疆
# ============================================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE, encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)
log = logging.getLogger(__name__)

# ============================================================
#  宸ュ叿鍑芥暟
# ============================================================

def get_mail_domain(email_addr):
    """浠庨偖绠卞湴鍧€鎻愬彇鍩熷悕锛屽 'user@qq.com' 鈫?'qq.com'"""
    m = re.search(r"@(.+)$", email_addr)
    return m.group(1).lower() if m else None


def get_imap_config(email_addr):
    domain = get_mail_domain(email_addr)
    if domain in IMAP_CONFIG:
        return IMAP_CONFIG[domain]
    # 灏濊瘯閫氱敤鎺ㄦ柇
    if "qq" in domain:
        return IMAP_CONFIG["qq.com"]
    if "163" in domain:
        return IMAP_CONFIG["163.com"]
    raise ValueError(f"涓嶆敮鎸佺殑閭鍩熷悕锛歿domain}锛岀洰鍓嶆敮鎸?QQ 閭鍜?163 閭銆?)


def get_smtp_config(email_addr):
    domain = get_mail_domain(email_addr)
    if domain in SMTP_CONFIG:
        return SMTP_CONFIG[domain]
    if "qq" in domain:
        return SMTP_CONFIG["qq.com"]
    if "163" in domain:
        return SMTP_CONFIG["163.com"]
    raise ValueError(f"涓嶆敮鎸佺殑閭鍩熷悕锛歿domain}锛岀洰鍓嶆敮鎸?QQ 閭鍜?163 閭銆?)


# ---- UID 杩借釜 ----
def load_processed_uids(filepath):
    if os.path.exists(filepath):
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                return set(json.load(f))
        except Exception:
            return set()
    return set()


def save_processed_uids(filepath, uid_set):
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(list(uid_set), f)


# ============================================================
#  閭欢瑙ｆ瀽
# ============================================================

def parse_email_body(msg):
    """
    浠?email.Message 涓彁鍙栫函鏂囨湰姝ｆ枃銆?    浼樺厛 text/plain锛屽叾娆?text/html锛堢畝鍗曞幓闄ゆ爣绛撅級銆?    """
    import email as em

    body_plain = ""
    body_html = ""

    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            cdisp = str(part.get("Content-Disposition", ""))
            if "attachment" in cdisp:
                continue
            payload = part.get_payload(decode=True)
            if not payload:
                continue
            charset = part.get_content_charset() or "utf-8"
            try:
                text = payload.decode(charset, errors="replace")
            except Exception:
                text = payload.decode("utf-8", errors="replace")
            if ctype == "text/plain":
                body_plain += text
            elif ctype == "text/html":
                body_html += text
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            charset = msg.get_content_charset() or "utf-8"
            try:
                body_plain = payload.decode(charset, errors="replace")
            except Exception:
                body_plain = payload.decode("utf-8", errors="replace")

    if body_plain:
        return body_plain.strip()
    if body_html:
        # 绠€鏄?HTML 鈫?鏂囨湰
        text = re.sub(r"<br\s*/?>", "\n", body_html, flags=re.IGNORECASE)
        text = re.sub(r"<[^>]+>", "", text)
        text = re.sub(r"&nbsp;", " ", text)
        text = re.sub(r"&lt;", "<", text)
        text = re.sub(r"&gt;", ">", text)
        text = re.sub(r"&amp;", "&", text)
        return text.strip()
    return ""


def parse_csv_from_body(body_text):
    """
    浠庨偖浠舵鏂囦腑瑙ｆ瀽 CSV 琛ㄦ牸鏁版嵁銆?    鏀寔锛?      1. 绗竴琛屾槸琛ㄥご锛堜腑鏂囧垪鍚嶏級
      2. 鍚庣画姣忚鏄竴鏉℃暟鎹?      3. 鍒椾箣闂寸敤閫楀彿銆佸埗琛ㄧ鎴栦腑鏂囧叏瑙掗€楀彿鍒嗛殧
    杩斿洖 (headers, records) 鎴?(None, []) 濡傛灉瑙ｆ瀽澶辫触銆?    """
    if not body_text:
        return None, []

    lines = [l.strip() for l in body_text.splitlines() if l.strip()]
    if len(lines) < 2:
        return None, []

    # 璺宠繃闈炶〃鏍艰锛堝 #寮€澶淬€佺被鍨嬫爣璁扮瓑锛?    data_lines = []
    for l in lines:
        if l.startswith("#") or l.startswith("绫诲瀷") or l.startswith("妯℃澘"):
            continue
        data_lines.append(l)

    if len(data_lines) < 2:
        return None, []

    # 鑷姩妫€娴嬪垎闅旂锛氫紭鍏堥€楀彿锛屽叾娆?tab
    sample = data_lines[0]
    if "\t" in sample:
        sep = "\t"
    elif "锛? in sample:
        sep = "锛?
    elif "," in sample:
        sep = ","
    else:
        return None, []

    # 瑙ｆ瀽琛ㄥご
    headers = [h.strip() for h in data_lines[0].split(sep)]
    if not headers:
        return None, []

    # 瑙ｆ瀽鏁版嵁琛?    records = []
    for line in data_lines[1:]:
        vals = [v.strip() for v in line.split(sep)]
        if len(vals) < len(headers):
            vals += [""] * (len(headers) - len(vals))
        elif len(vals) > len(headers):
            vals = vals[:len(headers)]
        rec = {}
        for i, h in enumerate(headers):
            rec[h] = vals[i] if i < len(vals) else ""
        records.append(rec)

    return headers, records


def extract_excel_attachments(msg):
    """
    浠庨偖浠朵腑鎻愬彇 Excel 闄勪欢锛?xlsx / .xls锛夛紝淇濆瓨鍒颁复鏃舵枃浠躲€?    杩斿洖 [filepath, ...]
    """
    import email as em

    temp_files = []
    if not msg.is_multipart():
        return temp_files

    for part in msg.walk():
        cdisp = str(part.get("Content-Disposition", ""))
        if "attachment" not in cdisp:
            continue
        filename = part.get_filename()
        if not filename:
            continue
        # 瑙ｇ爜涓枃鏂囦欢鍚?        try:
            from email.header import decode_header
            dh = decode_header(filename)
            filename = "".join(
                t.decode(c or "utf-8", errors="replace") if isinstance(t, bytes) else str(t)
                for t, c in dh
            )
        except Exception:
            pass

        if not (filename.lower().endswith(".xlsx") or filename.lower().endswith(".xls")):
            continue

        payload = part.get_payload(decode=True)
        if not payload:
            continue

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
        tmp.write(payload)
        tmp.close()
        temp_files.append(tmp.name)
        log.info(f"鎻愬彇 Excel 闄勪欢锛歿filename} 鈫?{tmp.name}")

    return temp_files


def parse_excel_data(filepath):
    """
    浠?Excel 鏂囦欢璇诲彇鏁版嵁銆?    绗竴琛屼负琛ㄥご锛屽悗缁瘡琛屼负涓€鏉℃暟鎹€?    杩斿洖 (headers, records) 鎴?(None, [])銆?    """
    try:
        wb = load_workbook(filepath, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(min_row=1, values_only=True))
        wb.close()

        if len(rows) < 2:
            return None, []

        headers = [str(h).strip() if h else "" for h in rows[0]]
        if not any(headers):
            return None, []

        records = []
        for row in rows[1:]:
            if row[0] is None and all(v is None for v in row):
                continue
            rec = {}
            for i, h in enumerate(headers):
                val = row[i] if i < len(row) else ""
                if val is None:
                    val = ""
                rec[h] = str(val).strip()
            # 杩囨护锛氳烦杩囦俊鎭笉瀹屾暣鐨勮锛堝彧鏈夊簭鍙?澶囨敞绛夊皯浜?涓湁鏁堝瓧娈碉級
            skip_keywords = ["搴忓彿", "澶囨敞"]
            filled = sum(1 for k, v in rec.items() if v and not any(s in k for s in skip_keywords))
            if filled < 2:
                log.info(f"  璺宠繃涓嶅畬鏁磋: {str(rec)[:100]}")
                continue
            records.append(rec)

        wb.close()
        log.info(f"Excel 琛ㄥご: {headers}")
        if records:
            log.info(f"Excel 棣栬: {records[0]}")
        return headers, records

    except Exception as e:
        log.warning(f"瑙ｆ瀽 Excel 澶辫触锛歿e}")
        return None, []


def detect_doc_type(headers):
    """
    鏍规嵁琛ㄥご鑷姩鍒ゆ柇鏂囨。绫诲瀷銆?    鍑哄簱閫氱煡鍗曠壒寰佸垪锛氱洰鐨勬腐銆佽埞鍙枫€佽鍒掗噺
    鎻愯揣濮旀墭鍑界壒寰佸垪锛氬叕鍙稿悕銆佹彁璐ц鍒掗噺銆佽埞鑸惰仈绯讳汉
    """
    鍑哄簱閫氱煡鍗昣keys = ["鐩殑娓?, "鑸瑰彿", "璁″垝閲?]
    鎻愯揣濮旀墭鍑絖keys = ["鍏徃鍚?, "鎻愯揣璁″垝閲?, "鑸硅埗鑱旂郴浜?]

    score_a = sum(1 for k in 鍑哄簱閫氱煡鍗昣keys if any(k in h for h in headers))
    score_b = sum(1 for k in 鎻愯揣濮旀墭鍑絖keys if any(k in h for h in headers))

    if score_b > score_a:
        return "鎻愯揣濮旀墭鍑?
    if score_a > 0:
        return "鍑哄簱閫氱煡鍗?
    # fallback: 鏍规嵁鍒楁暟鎺ㄦ柇
    if len(headers) <= 5:
        return "鍑哄簱閫氱煡鍗?
    return "鎻愯揣濮旀墭鍑?


# ============================================================
#  Word 鐢熸垚 鈥?鍑哄簱閫氱煡鍗?# ============================================================

def _get_field(rec, *keys):
    """浠?record 涓彇瀛楁锛屾敮鎸佸涓€欓€?key 鍜岄儴鍒嗗尮閰?""
    for k in keys:
        if k in rec:
            return rec[k]
    # 閮ㄥ垎鍖归厤锛歬ey 鍖呭惈鍦ㄨ〃澶翠腑
    for k in keys:
        for h in rec:
            if k in h:
                return rec[h]
    return ""


def generate_chukutongzhidan(records, date_chinese, date_num):
    """
    鐢熸垚銆屾补鍝佽鑸瑰嚭搴撻€氱煡鍗曘€峎ord 鏂囨。銆?    鎵€鏈?records 鏀惧湪鍚屼竴涓〃鏍间腑锛屾瘡琛屼竴涓嚭搴撶紪鍙枫€?    杩斿洖涓存椂鏂囦欢璺緞銆?    """
    doc = Document()

    # 榛樿椤佃竟璺?    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "浠垮畫"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "浠垮畫")

    # 鏍囬
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("涓婃捣涔呰仈闆嗗洟鏈夐檺鍏徃")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "浠垮畫"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "浠垮畫")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("娌瑰搧瑁呰埞鍑哄簱閫氱煡鍗?)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "浠垮畫"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "浠垮畫")

    # 鍙戜欢浜轰俊鎭?    doc.add_paragraph(
        f"鍙戜欢浜猴細瀹嬪崡甯?       鐢佃瘽锛?8018586154                 鏃ユ湡锛歿date_chinese}"
    )

    doc.add_paragraph(" 杩戞湡灏嗙敱浠ヤ笅鑸硅埗鍒版补搴撹娌癸紝璇锋彁鍓嶅畨鎺掕埞鍚嶇敵鎶ャ€佽鑸瑰嚭搴撶瓑鐩稿叧浜嬪疁銆?)

    doc.add_paragraph()  # 绌鸿

    # 琛ㄦ牸
    headers = ["鍑哄簱缂栧彿", "鐩殑娓?, "鑸瑰彿", "娌瑰搧鍝佸彿", "浣滀笟缃愬彿", "璁″垝閲忥紙鍚級", "澶囨敞", "鍏佽瑁呰揣鏃堕棿锛堝惈锛?]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 琛ㄥご
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "浠垮畫"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "浠垮畫")

    # 鏁版嵁琛?    for idx, rec in enumerate(records):
        counter = 99 + idx
        qty = _safe_float(_get_field(rec, "璁″垝閲?, "鎻愯揣璁″垝閲?))
        code = f"ZX-C0-0999-{date_num}-{qty:.2f}-{counter:03d}"

        row = table.add_row()
        values = [
            code,
            _get_field(rec, "鐩殑娓?),
            _get_field(rec, "鑸瑰彿", "鑸瑰悕"),
            "0鍙疯溅鐢ㄦ煷娌癸紙VI锛?,
            "T107/109",
            f"{qty:.2f}",
            _get_field(rec, "澶囨敞"),
            _get_field(rec, "鍏佽瑁呰揣鏃堕棿"),
        ]
        for i, v in enumerate(values):
            row.cells[i].text = v
            for para in row.cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "浠垮畫"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "浠垮畫")

    doc.add_paragraph()
    doc.add_paragraph(f"鍒惰〃锛氬畫鍗楀笇                              澶嶆牳锛?)

    # 淇濆瓨鍒颁复鏃舵枃浠?    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


# ============================================================
#  Word 鐢熸垚 鈥?鎻愯揣濮旀墭鍑?# ============================================================

def generate_tihuoweituohan(rec, date_chinese):
    """
    鐢熸垚銆屾彁璐у鎵樺嚱銆峎ord 鏂囨。锛堝崟鏉¤褰曪級銆?    杩斿洖涓存椂鏂囦欢璺緞銆?    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "浠垮畫"
    font.size = Pt(14)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "浠垮畫")

    def add_para(text, bold=False, size=14, align="left", font_name="浠垮畫"):
        p = doc.add_paragraph()
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        return p

    qty = _safe_float(_get_field(rec, "鎻愯揣璁″垝閲?, "璁″垝閲?))
    company = _get_field(rec, "鍏徃鍚?) or "涓婃捣涔呰仈闆嗗洟鐭虫补鍖栧伐鏈夐檺鍏徃"
    ship = _get_field(rec, "鑸瑰悕", "鑸瑰彿")
    dest = _get_field(rec, "鐩殑娓?)
    contact = _get_field(rec, "鑸硅埗鑱旂郴浜?)
    phone = _get_field(rec, "鑱旂郴鐢佃瘽")
    eta = _get_field(rec, "棰勮鍒版腐鏃堕棿")
    remark = _get_field(rec, "澶囨敞")

    add_para("鎻愯揣濮旀墭鍑?, bold=True, size=16, align="center")
    doc.add_paragraph()
    add_para(f"{company}锛?)
    add_para(
        f"    鎴戝叕鍙稿鎵樿埞鍙埌璐靛叕鍙告寚瀹氱爜澶存彁鍙?鍙疯溅鐢ㄦ煷娌癸紙VI锛墈qty:.2f}鍚紝鍏蜂綋鑸硅埗淇℃伅濡備笅锛?
    )

    # 淇℃伅琛ㄦ牸
    table = doc.add_table(rows=9, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    def set_cell(row, col, text, width=None, span=1, bold=False):
        cell = table.rows[row].cells[col]
        cell.text = text
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "浠垮畫"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "浠垮畫")
                run.bold = bold
        # 鍚堝苟鍗曞厓鏍?        if span > 1:
            for j in range(1, span):
                cell.merge(table.rows[row].cells[col + j])
        return cell

    # 濉厖琛ㄦ牸鏁版嵁锛堟ā鏉垮浐瀹氱粨鏋勶級
    set_cell(0, 0, "鑸瑰悕", bold=True);          set_cell(0, 1, ship, span=3)
    set_cell(1, 0, "鍝佸悕", bold=True);          set_cell(1, 1, "0鍙疯溅鐢ㄦ煷娌癸紙VI锛?, span=3)
    set_cell(2, 0, "鎻愯揣璁″垝閲?, bold=True);     set_cell(2, 1, f"{qty:.2f}鍚?, span=3)
    set_cell(3, 0, "鐩殑娓?, bold=True);         set_cell(3, 1, dest, span=3)
    set_cell(4, 0, "涓嬩竴娓?, bold=True);         set_cell(4, 1, "涓婃捣鐭虫礊鍙?, span=3)
    set_cell(5, 0, "鑸硅埗鑱旂郴浜?, bold=True);     set_cell(5, 1, contact); set_cell(5, 2, "鑱旂郴鐢佃瘽", bold=True); set_cell(5, 3, phone)
    set_cell(6, 0, "涓氬姟鑱旂郴浜?, bold=True);     set_cell(6, 1, "榫氳€€鍏?); set_cell(6, 2, "鑱旂郴鐢佃瘽", bold=True); set_cell(6, 3, "13601650745")
    set_cell(7, 0, "棰勮鍒版腐鏃堕棿", bold=True);   set_cell(7, 1, eta, span=3)
    set_cell(8, 0, "澶囨敞", bold=True);           set_cell(8, 1, remark, span=3)

    doc.add_paragraph()
    add_para("    璇疯吹鍏徃浜堜互瀹夋帓瑁呰揣璁″垝銆?)
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_para("涓婃捣涔呰仈闆嗗洟鏈夐檺鍏徃", align="right")
    add_para(date_chinese, align="right")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


def _cleanup_temps(file_list):
    for f in file_list:
        try:
            os.remove(f)
        except Exception:
            pass


def _safe_float(val):
    try:
        return float(val)
    except (ValueError, TypeError):
        return 0.0


# ============================================================
#  閭欢鎿嶄綔锛圛MAP + SMTP锛?# ============================================================

def connect_imap():
    """杩炴帴鍒?IMAP 鏀朵欢鏈嶅姟鍣紝杩斿洖 imaplib.IMAP4_SSL 瀵硅薄"""
    import imaplib
    cfg = get_imap_config(ACCOUNT["email"])
    log.info(f"IMAP 杩炴帴锛歿cfg['host']}:{cfg['port']}")
    conn = imaplib.IMAP4_SSL(cfg["host"], cfg["port"])
    conn.login(ACCOUNT["email"], ACCOUNT["password"])
    log.info("IMAP 鐧诲綍鎴愬姛")
    return conn


def fetch_new_emails(imap_conn, processed_uids):
    """
    浠?INBOX 鎷夊彇鏈鐞嗙殑鏂伴偖浠躲€?    QQ IMAP 鐨?SUBJECT SEARCH 涓嶅噯纭紝鐩存帴鎷夊彇鏈€杩戦偖浠跺湪 Python 渚ц繃婊や富棰樸€?    杩斿洖 [(msg_uid, email.Message), ...]
    """
    import email as em

    imap_conn.select("INBOX", readonly=False)

    # 鎷夊彇鏈€杩?100 灏侀偖浠讹紙涓嶄緷璧?IMAP 鐨?SUBJECT 鎼滅储锛?    status, data = imap_conn.uid("SEARCH", None, "ALL")
    if status != "OK" or not data[0]:
        log.info("鏈壘鍒颁换浣曢偖浠?)
        return []

    all_uids = sorted(int(u) for u in data[0].split())
    recent_uids = all_uids[-5:]

    log.info(f"鏀朵欢绠卞叡 {len(all_uids)} 灏侊紝鎵弿鏈€杩?{len(recent_uids)} 灏?)

    results = []
    for uid in recent_uids:
        if uid in processed_uids:
            continue
        try:
            # 鐩存帴鑾峰彇瀹屾暣閭欢
            status, msg_data = imap_conn.uid("FETCH", str(uid).encode(), "(RFC822)")
            if status != "OK" or not msg_data[0]:
                continue
            raw_email = None
            for item in msg_data:
                if isinstance(item, tuple) and len(item) >= 2:
                    raw_email = item[1]
                    break
            if not raw_email:
                continue

            msg = em.message_from_bytes(raw_email, policy=email.policy.default)

            # 鏉′欢1锛氫富棰樺繀椤讳负绌?            subject = str(msg.get("Subject", "")).strip()
            if subject:
                continue

            # 鏉′欢2锛氬繀椤绘湁 Excel 闄勪欢锛?xlsx / .xls锛?            has_excel = False
            if msg.is_multipart():
                for part in msg.walk():
                    cdisp = str(part.get("Content-Disposition", ""))
                    if "attachment" not in cdisp:
                        continue
                    fname = part.get_filename() or ""
                    if fname.lower().endswith((".xlsx", ".xls")):
                        has_excel = True
                        break

            if not has_excel:
                continue

            log.info(f"  [{uid}] 绌轰富棰?+ Excel闄勪欢")
            results.append((uid, msg))

            if MARK_AS_READ:
                imap_conn.uid("STORE", str(uid).encode(), "+FLAGS", "\\Seen")
        except Exception as e:
            log.warning(f"鑾峰彇閭欢 UID={uid} 澶辫触锛歿e}")

    log.info(f"鎵弿瀹屾垚锛屽尮閰?{len(results)} 灏?)
    return results


def send_reply(original_msg, attachments, body_text=""):
    """
    鍥炲閭欢锛屽甫闄勪欢銆?    original_msg: 鍘熷 email.Message锛堣幏鍙栧彂浠朵汉銆佷富棰橈級
    attachments: [(filename, filepath), ...]
    """
    import smtplib
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders
    from email.utils import formataddr

    sender = ACCOUNT["email"]
    to_addr = original_msg["From"]
    orig_subject = original_msg.get("Subject", "")

    # 鏋勫缓鍥炲涓婚
    if orig_subject.startswith("Re:"):
        reply_subject = orig_subject
    else:
        reply_subject = f"Re: {orig_subject}"

    # 鏋勫缓閭欢
    msg = MIMEMultipart()
    msg["From"] = sender
    msg["To"] = to_addr
    msg["Subject"] = reply_subject
    msg["In-Reply-To"] = original_msg.get("Message-ID", "")
    msg["References"] = original_msg.get("Message-ID", "")

    if not body_text:
        body_text = f"鎮ㄥソ锛孿n\n宸叉牴鎹偍鐨勯偖浠剁敓鎴?Word 鏂囨。锛岃瑙侀檮浠躲€俓n\n鐢熸垚鏃堕棿锛歿datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n姝ら偖浠剁敱绯荤粺鑷姩鍙戝嚭銆?

    msg.attach(MIMEText(body_text, "plain", "utf-8"))

    for fname, fpath in attachments:
        with open(fpath, "rb") as f:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f'attachment; filename="{fname}"',
            )
            msg.attach(part)

    # 鍙戦€?    smtp_cfg = get_smtp_config(ACCOUNT["email"])
    log.info(f"SMTP 杩炴帴锛歿smtp_cfg['host']}:{smtp_cfg['port']}")
    with smtplib.SMTP_SSL(smtp_cfg["host"], smtp_cfg["port"]) as smtp:
        smtp.login(ACCOUNT["email"], ACCOUNT["password"])
        smtp.send_message(msg)

    log.info(f"鍥炲宸插彂閫?鈫?{to_addr}")


# ============================================================
#  鏍稿績澶勭悊閫昏緫
# ============================================================

def process_email(uid, msg):
    """
    澶勭悊鍗曞皝閭欢锛?      1. 瑙ｆ瀽姝ｆ枃
      2. 鎻愬彇琛ㄦ牸鏁版嵁
      3. 鐢熸垚 Word
      4. 鍥炲闄勪欢
    杩斿洖 True 琛ㄧず澶勭悊鎴愬姛銆?    """
    sender = msg["From"]
    subject = msg.get("Subject", "")
    log.info(f"澶勭悊閭欢 UID={uid} | 鍙戜欢浜猴細{sender} | 涓婚锛歿subject}")

    headers, records = None, []
    excel_temps = []

    # 1) 浼樺厛灏濊瘯 Excel 闄勪欢
    excel_files = extract_excel_attachments(msg)
    for xf in excel_files:
        h, r = parse_excel_data(xf)
        if h and r:
            headers = h
            records = r
            excel_temps.append(xf)
            break  # 鍙鐞嗙涓€涓湁鏁?Excel 闄勪欢
        else:
            excel_temps.append(xf)

    # 2) 鑻ユ病鏈?Excel 闄勪欢锛屽垯灏濊瘯瑙ｆ瀽閭欢姝ｆ枃 CSV
    if not records:
        body = parse_email_body(msg)
        log.debug(f"閭欢姝ｆ枃锛堝墠500瀛楋級锛歿body[:500] if body else '(绌?'}")
        if body:
            headers, records = parse_csv_from_body(body)

    if not records:
        log.warning(f"閭欢 UID={uid} 鏈瘑鍒埌鏁版嵁锛岃烦杩?)
        send_reply(
            msg, [],
            "鎮ㄥソ锛孿n\n鏈兘浠庢偍鐨勯偖浠朵腑璇嗗埆鍒版暟鎹€俓n\n"
            "鏀寔涓ょ鏂瑰紡锛歕n"
            "  1. 鍦ㄩ偖浠朵腑**娣诲姞 Excel 闄勪欢**锛?xlsx锛夛紝绗竴琛屼负琛ㄥご锛屽悗缁负鏁版嵁琛孿n"
            "  2. 鍦ㄩ偖浠舵鏂囦腑绮樿创 CSV 鏁版嵁锛屾牸寮忓涓嬶細\n\n"
            "銆愬嚭搴撻€氱煡鍗曘€慭n"
            "鐩殑娓?鑸瑰彿,璁″垝閲?澶囨敞,鍏佽瑁呰揣鏃堕棿\n"
            "涓婃捣娓?姹熸樊娌?98,476.50,鍛ㄤ簹骞?13524901555,2026-06-01 08:00\n\n"
            "銆愭彁璐у鎵樺嚱銆慭n"
            "鍏徃鍚?鑸瑰悕,鎻愯揣璁″垝閲?鐩殑娓?鑸硅埗鑱旂郴浜?鑱旂郴鐢佃瘽,棰勮鍒版腐鏃堕棿,澶囨敞\n"
            "涓婃捣涔呰仈闆嗗洟鐭虫补鍖栧伐鏈夐檺鍏徃,姹熸樊娌?98,476.50,涓婃捣娓?鍛ㄤ簹骞?13524901555,2026-06-01 08:00,"
        )
        _cleanup_temps(excel_temps)
        return False

    doc_type = detect_doc_type(headers)
    log.info(f"璇嗗埆鏂囨。绫诲瀷锛歿doc_type} | {len(records)} 鏉℃暟鎹?)

    today = datetime.now()
    date_chinese = today.strftime("%Y骞?m鏈?d鏃?)
    date_num = today.strftime("%Y%m%d")

    attachments = []
    temp_files = []

    try:
        if doc_type == "鍑哄簱閫氱煡鍗?:
            tmp = generate_chukutongzhidan(records, date_chinese, date_num)
            attachments.append((f"鍑哄簱閫氱煡鍗昣{date_num}.docx", tmp))
            temp_files.append(tmp)

        elif doc_type == "鎻愯揣濮旀墭鍑?:
            if len(records) == 1:
                tmp = generate_tihuoweituohan(records[0], date_chinese)
                ship = records[0].get("鑸瑰悕", records[0].get("鑸瑰彿", "unknown"))
                fname = f"鎻愯揣濮旀墭鍑絖{ship}.docx"
                attachments.append((fname, tmp))
                temp_files.append(tmp)
            else:
                # 澶氫唤鏂囨。鎵撳寘鎴?ZIP
                zip_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
                zip_tmp.close()
                temp_files.append(zip_tmp.name)
                with zipfile.ZipFile(zip_tmp.name, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, rec in enumerate(records):
                        tmp = generate_tihuoweituohan(rec, date_chinese)
                        temp_files.append(tmp)
                        ship = _get_field(rec, "鑸瑰悕", "鑸瑰彿") or f"record_{i+1}"
                        zf.write(tmp, f"鎻愯揣濮旀墭鍑絖{ship}.docx")
                attachments.append(("鎻愯揣濮旀墭鍑絖鎵归噺.zip", zip_tmp.name))

        send_reply(msg, attachments)
        return True

    except Exception as e:
        log.error(f"澶勭悊閭欢 UID={uid} 鍑洪敊锛歿e}", exc_info=True)
        try:
            send_reply(msg, [], f"鎮ㄥソ锛孿n\n鐢熸垚 Word 鏃跺嚭閿欙細{e}\n\n璇锋鏌ユ暟鎹牸寮忓悗閲嶈瘯銆?)
        except Exception:
            pass
        return False

    finally:
        # 娓呯悊涓存椂鏂囦欢
        for tmp in temp_files:
            try:
                os.remove(tmp)
            except Exception:
                pass
        _cleanup_temps(excel_temps)


# ============================================================
#  涓诲惊鐜?# ============================================================

def main_loop():
    """涓荤洃鎺у惊鐜?""
    print("=" * 56)
    print("  閭欢鐩戞帶 & 鑷姩鐢熸垚 Word 宸ュ叿")
    print("=" * 56)
    print(f"  鐩戞帶閭锛歿ACCOUNT['email']}")
    print(f"  妫€鏌ラ棿闅旓細{CHECK_INTERVAL} 绉?)
    print(f"  涓婚鍏抽敭璇嶏細{', '.join(SUBJECT_KEYWORDS)}")
    print(f"  鏃ュ織鏂囦欢锛歿LOG_FILE}")
    print("=" * 56)
    print("  鎸?Ctrl+C 鍋滄鐩戞帶")
    print()

    processed = load_processed_uids(TRACK_FILE)
    log.info(f"宸插姞杞?{len(processed)} 鏉″凡澶勭悊閭欢 UID")

    while True:
        try:
            imap = connect_imap()
            try:
                new_emails = fetch_new_emails(imap, processed)
                if new_emails:
                    log.info(f"鍙戠幇 {len(new_emails)} 灏佹柊閭欢锛屽紑濮嬪鐞?..")
                    for uid, msg in new_emails:
                        ok = process_email(uid, msg)
                        if ok:
                            processed.add(uid)
                        time.sleep(2)  # 澶勭悊闂撮殧锛岄伩鍏嶈闄愭祦
                    save_processed_uids(TRACK_FILE, processed)
                else:
                    log.info("鏃犳柊閭欢")
            finally:
                try:
                    imap.logout()
                except Exception:
                    pass

        except Exception as e:
            log.error(f"妫€鏌ラ偖浠舵椂鍑洪敊锛歿e}", exc_info=True)

        time.sleep(CHECK_INTERVAL)


def run_once():
    """浠呮鏌ヤ竴娆″苟澶勭悊锛岀敤浜庢祴璇?""
    print("=" * 56)
    print("  閭欢鐩戞帶 鈥?鍗曟妫€鏌ユā寮?)
    print("=" * 56)

    processed = load_processed_uids(TRACK_FILE)
    log.info(f"宸插姞杞?{len(processed)} 鏉″凡澶勭悊閭欢 UID")

    imap = connect_imap()
    try:
        new_emails = fetch_new_emails(imap, processed)
        if not new_emails:
            print("鏃犳柊閭欢闇€瑕佸鐞嗐€?)
            return

        print(f"鍙戠幇 {len(new_emails)} 灏佹柊閭欢锛屽紑濮嬪鐞?..")
        for uid, msg in new_emails:
            ok = process_email(uid, msg)
            if ok:
                processed.add(uid)
            time.sleep(1)

        save_processed_uids(TRACK_FILE, processed)
        print("澶勭悊瀹屾垚銆?)
    finally:
        try:
            imap.logout()
        except Exception:
            pass


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="閭欢鐩戞帶 & 鑷姩鐢熸垚 Word")
    parser.add_argument(
        "--once", action="store_true",
        help="鍙鏌ヤ竴娆″氨閫€鍑猴紙鐢ㄤ簬娴嬭瘯锛?
    )
    args = parser.parse_args()

    if args.once:
        run_once()
    else:
        main_loop()
